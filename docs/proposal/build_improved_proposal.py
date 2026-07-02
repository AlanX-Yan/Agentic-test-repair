from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Goal-Driven_Test_Repair_Proposal_Improved.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width.inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def style_table(table, widths=None):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_width(table)
    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                set_cell_width(cell, widths[idx])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(3)
                paragraph.paragraph_format.line_spacing = 1.12
                for run in paragraph.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(10)
            if row_idx == 0:
                set_cell_shading(cell, "F4F6F9")
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.208
        p.add_run(item)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = [Inches(1.65), Inches(4.85)]
    header = table.rows[0].cells
    header[0].text = "Item"
    header[1].text = "Proposed Plan"
    for left, right in rows:
        cells = table.add_row().cells
        cells[0].text = left
        cells[1].text = right
    style_table(table, widths)
    return table


def add_simple_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    style_table(table, widths)
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Calibri"
        styles[name].font.bold = True
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    styles["Heading 1"].paragraph_format.space_before = Pt(18)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    styles["Heading 2"].paragraph_format.space_before = Pt(12)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)
    styles["Heading 3"].paragraph_format.space_before = Pt(8)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    for name in ["List Bullet", "List Number"]:
        styles[name].font.name = "Calibri"
        styles[name].font.size = Pt(11)
        styles[name].paragraph_format.space_after = Pt(4)
        styles[name].paragraph_format.line_spacing = 1.208


def main():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Project Proposal")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    sr = subtitle.add_run(
        "Goal-Driven Repair of Agent-Generated Java Tests Using AromaDr Test Smell Feedback"
    )
    sr.font.name = "Calibri"
    sr.font.size = Pt(15)
    sr.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta.add_run("Course: ").bold = True
    meta.add_run("CSCI 4970W Advanced Project Laboratory\n")
    meta.add_run("Project type: ").bold = True
    meta.add_run("Research prototype and empirical software engineering evaluation\n")
    meta.add_run("Primary implementation target: ").bold = True
    meta.add_run("Java programs and JUnit tests, using AromaDr for test smell detection")

    add_heading(doc, "Executive Summary", 1)
    doc.add_paragraph(
        "This project proposes a goal-driven repair framework for improving Java "
        "tests generated by AI coding agents. A coding agent first generates JUnit "
        "tests for a target Java method, class, bug fix, or pull request. Instead "
        "of accepting the first test suite, a master agent evaluates the tests "
        "against explicit quality criteria, including whether the tests compile, "
        "whether they pass, whether assertions are meaningful, whether mocks are "
        "appropriate, whether boundary and error cases are covered, and whether "
        "AromaDr reports detectable test smells. When the criteria are not "
        "satisfied, the master agent returns targeted smell-based feedback to the "
        "coding agent and asks it to repair the tests. The loop continues until "
        "the criteria are met or a fixed iteration budget is reached."
    )
    doc.add_paragraph(
        "The expected result is a working prototype and a small empirical study "
        "comparing one-shot agent-generated Java tests with tests repaired through "
        "AromaDr-guided feedback. The project is intended to be more than a toy "
        "test generator: it combines Java build/test execution, external test "
        "smell analysis, agent orchestration, quality criteria, and before-after "
        "evaluation."
    )

    add_heading(doc, "Motivation and Research Gap", 1)
    doc.add_paragraph(
        "Coding agents such as Codex, Claude Code, Copilot, Cursor, and similar "
        "systems are increasingly able to produce unit tests. In Java projects, "
        "however, a generated JUnit test can compile and pass while still providing "
        "weak behavioral evidence. For example, a test may exercise a method but "
        "use overly broad assertions, duplicate setup logic, depend on unnecessary "
        "mocks, or miss boundary and exception cases. These issues are especially "
        "important for test-smell research because generated tests may increase "
        "quantity without improving maintainability or fault-detection value."
    )
    doc.add_paragraph(
        "Recent empirical work by Milanese, Salzano, Spina, Vitale, Pareschi, "
        "Fasano, and Fazzini characterizes testing in human-agent pull requests "
        "and compares it with human pull requests using dimensions such as testing "
        "extent, test evolution tasks, and test smells. This project builds on "
        "that direction but changes the question from characterization to "
        "intervention: if coding agents tend to create or expand tests, can a "
        "goal-driven feedback loop make those tests better before they are "
        "accepted?"
    )
    doc.add_paragraph(
        "Following Professor Mattia Fazzini's suggestion, this revised proposal "
        "narrows the implementation target to Java programs and proposes using "
        "AromaDr: A Language-Independent Tool for Detecting Test Smells as the "
        "primary smell detector. This makes the project better aligned with the "
        "test-smell domain and avoids spending the semester re-implementing a "
        "detector from scratch."
    )

    add_heading(doc, "Project Objectives", 1)
    add_bullets(
        doc,
        [
            "Design a goal-driven repair loop in which a master agent evaluates generated tests against explicit quality criteria.",
            "Implement a prototype pipeline for Java projects that generates, compiles, executes, analyzes, and repairs JUnit tests.",
            "Integrate AromaDr as the primary tool for detecting test smells in generated and repaired tests.",
            "Translate AromaDr reports and build/test results into actionable feedback that a coding agent can use to revise existing tests.",
            "Evaluate whether repaired tests improve over the original generated tests using measurable before-after metrics.",
        ],
    )

    add_heading(doc, "Research Questions", 1)
    add_simple_table(
        doc,
        ["RQ", "Question", "Evidence Collected"],
        [
            (
                "RQ1",
                "What quality issues and test smells commonly appear in coding-agent-generated Java tests?",
                "AromaDr smell categories, JUnit assertion patterns, mock usage, compilation failures, test failures, and missing-case labels from the initial generated tests.",
            ),
            (
                "RQ2",
                "Does test-smell feedback help coding agents repair low-quality generated tests?",
                "Before-after changes in AromaDr smell count, assertion quality, mock usage, compilation success, JUnit pass rate, and coverage.",
            ),
            (
                "RQ3",
                "Are repaired tests more behaviorally useful than the original generated tests?",
                "Assertion specificity, changed-line or method coverage, edge/error case coverage, and optional mutation score from a Java mutation-testing tool.",
            ),
        ],
        [Inches(0.65), Inches(3.25), Inches(2.6)],
    )

    add_heading(doc, "System Design", 1)
    doc.add_paragraph(
        "The system has three main roles: a coding agent that writes and revises "
        "tests, a detector that measures objective quality signals, and a master "
        "agent that decides whether the current tests satisfy the criteria."
    )
    add_simple_table(
        doc,
        ["Component", "Responsibility"],
        [
            (
                "Coding agent",
                "Generates the first JUnit test file and later repairs that file using feedback from the master agent.",
            ),
            (
                "Execution harness",
                "Runs Maven or Gradle test commands, records compilation and test failures, and optionally measures coverage using JaCoCo.",
            ),
            (
                "AromaDr integration",
                "Runs AromaDr on generated and repaired Java/JUnit tests, normalizes its output, and stores smell counts and locations for comparison.",
            ),
            (
                "Master agent",
                "Combines execution results, smell reports, and optional LLM-based judgment to decide whether the tests meet the project criteria.",
            ),
            (
                "Feedback generator",
                "Converts AromaDr findings and execution failures into concrete repair instructions, such as splitting broad tests, improving assertions, or adding missing exception cases.",
            ),
        ],
        [Inches(1.55), Inches(4.95)],
    )

    add_heading(doc, "Goal-Driven Loop", 2)
    add_numbered(
        doc,
        [
            "Input a target Java method, class, code change, or PR diff.",
            "Ask the coding agent to generate JUnit tests for the target behavior.",
            "Run the generated tests and collect compilation, pass/fail, and coverage information.",
            "Run AromaDr on the generated test files and normalize the smell report.",
            "Have the master agent evaluate the tests against explicit criteria.",
            "If the criteria are not met, produce targeted AromaDr-based feedback and ask the coding agent to repair the same tests.",
            "Repeat until the criteria are satisfied or the maximum repair iteration count is reached.",
        ],
    )
    doc.add_paragraph(
        "The key design choice is that the coding agent is not allowed to declare "
        "success by itself. The master agent independently checks the tests against "
        "criteria, which matches the goal-driven process described in the provided "
        "Goal-Driven README."
    )

    add_heading(doc, "Initial Test Smell Scope", 1)
    add_simple_table(
        doc,
        ["Smell or Quality Issue", "Detection Strategy", "Repair Feedback Example"],
        [
            (
                "AromaDr-reported smells",
                "Run AromaDr on each generated and repaired Java test suite; use its reported smell type, file, and location as the primary signal.",
                "Address the reported smell directly, for example by simplifying setup, splitting tests, or replacing fragile patterns.",
            ),
            (
                "Missing or weak assertion",
                "Supplement AromaDr with lightweight JUnit checks for tests that execute code but do not verify concrete outcomes.",
                "Add explicit assertions for return values, state changes, thrown exceptions, or object fields.",
            ),
            (
                "Excessive or unnecessary mocking",
                "Count Mockito usage and compare with dependency boundaries; flag tests that mainly verify mocks rather than behavior.",
                "Use real collaborators when practical and reserve mocks for external, slow, nondeterministic, or hard-to-control dependencies.",
            ),
            (
                "Generic test name",
                "Name pattern rule for JUnit methods such as test1, testBasic, or names that do not describe behavior.",
                "Rename the test to describe the behavior and expected outcome.",
            ),
            (
                "Missing edge/error cases",
                "Heuristic check using method signature, branches, exceptions, nullability, and boundary constants.",
                "Add tests for null, empty, boundary, invalid, and documented exception behavior where relevant.",
            ),
            (
                "Overly broad test",
                "Use AromaDr findings plus method length, assertion count, and multiple unrelated calls in one test.",
                "Split the test into focused cases with one primary behavior per test.",
            ),
        ],
        [Inches(1.45), Inches(2.55), Inches(2.5)],
    )

    add_heading(doc, "Success Criteria", 1)
    doc.add_paragraph(
        "The prototype will use criteria that are concrete enough for the master "
        "agent to evaluate automatically or semi-automatically."
    )
    add_bullets(
        doc,
        [
            "All generated Java tests compile and run under the project's Maven or Gradle test command.",
            "Every JUnit test contains at least one meaningful assertion or an explicit exception check.",
            "AromaDr-reported smell counts are reduced after repair, especially for repeated or high-confidence smell categories.",
            "Mockito or other mock usage remains below a configurable threshold unless mocks are justified by external dependencies.",
            "JaCoCo coverage does not decrease after repair and ideally improves for branches, boundary cases, or changed lines.",
            "The repaired test suite includes at least one normal case plus relevant edge, null, invalid, or exception cases when such cases can be inferred.",
            "The repair loop terminates after meeting criteria or after a fixed budget, such as three repair iterations.",
        ],
    )

    add_heading(doc, "Evaluation Plan", 1)
    doc.add_paragraph(
        "The evaluation will compare the initial tests produced by the coding "
        "agent with the repaired tests produced after the goal-driven feedback "
        "loop. The main analysis will use paired before-after comparisons over "
        "the same target functions or changes."
    )
    add_simple_table(
        doc,
        ["Metric", "Purpose"],
        [
            ("Test pass rate", "Shows whether repair keeps tests executable and integrated with the target code."),
            ("AromaDr smell count and smell delta", "Measures whether smell-based feedback reduces detected test smells."),
            ("Assertion count and specificity", "Captures whether JUnit tests verify concrete behavior rather than only executing code."),
            ("Mockito/mock usage", "Checks whether repaired tests avoid testing mocks instead of real business logic."),
            ("JaCoCo coverage", "Verifies that repair does not reduce statement, branch, or changed-line coverage."),
            ("Mutation score, stretch", "Estimates whether tests catch behavioral faults, using PIT or a similar Java mutation-testing tool if time allows."),
        ],
        [Inches(2.2), Inches(4.3)],
    )

    add_heading(doc, "Benchmark and Scope", 1)
    doc.add_paragraph(
        "The revised implementation will focus on Java and JUnit because this "
        "aligns with the current test-smell research direction and with Professor "
        "Fazzini's recommendation to use AromaDr. A feasible benchmark for one "
        "semester is 20 to 40 test-generation tasks drawn from small Java classes, "
        "open-source methods, bug-fix examples, or benchmark-style projects with "
        "known behavior. The project can begin with Maven-based Java projects and "
        "then extend to Gradle projects or small PR diffs if the core loop is "
        "stable."
    )
    add_simple_table(
        doc,
        ["Version", "Scope", "Why It Matters"],
        [
            (
                "Minimum viable version",
                "Java methods/classes, JUnit generation, Maven test execution, AromaDr smell reports, one to three repair iterations, 20 tasks.",
                "Demonstrates the core research idea with a reliable prototype.",
            ),
            (
                "Strong semester version",
                "Adds JaCoCo coverage, branch or changed-line coverage, better edge-case inference, 30 to 40 tasks, and manual spot-checking.",
                "Makes the evaluation stronger and resume-ready.",
            ),
            (
                "Stretch version",
                "Adds PIT mutation testing, Gradle support, or real GitHub PR diffs.",
                "Shows whether repaired tests catch behavioral faults, not only whether they look cleaner.",
            ),
        ],
        [Inches(1.35), Inches(3.1), Inches(2.05)],
    )

    add_heading(doc, "Risks and Mitigations", 1)
    add_simple_table(
        doc,
        ["Risk", "Mitigation"],
        [
            (
                "LLM feedback may be subjective or inconsistent.",
                "Use AromaDr and deterministic build/test results for primary metrics and reserve LLM judgment for semantic cases that tools cannot capture.",
            ),
            (
                "Repair may overfit to detector rules.",
                "Track coverage, assertion specificity, pass rate, and optional mutation score so the system is not rewarded only for hiding AromaDr smells.",
            ),
            (
                "Java project setup may be harder than small scripted examples.",
                "Start with Maven projects and controlled Java classes, then add selected open-source projects once the harness is stable.",
            ),
            (
                "Mutation testing may exceed the course time budget.",
                "Keep mutation score as a stretch metric; the core evaluation can stand on smell delta, pass rate, assertions, and coverage.",
            ),
        ],
        [Inches(2.0), Inches(4.5)],
    )

    add_heading(doc, "Timeline", 1)
    add_simple_table(
        doc,
        ["Weeks", "Milestone"],
        [
            ("1-2", "Review related work, study AromaDr input/output requirements, define success criteria, select initial Java benchmark tasks."),
            ("3-4", "Build JUnit generation and Maven test execution harness; log prompts, outputs, compilation results, and test results."),
            ("5-6", "Integrate AromaDr, normalize smell reports, and add lightweight JUnit-specific checks for assertions and mock usage."),
            ("7-8", "Implement master-agent feedback generation and repair loop using AromaDr findings and build/test results."),
            ("9-10", "Run pilot experiments on Java tasks, debug project setup and feedback failure modes, refine criteria and templates."),
            ("11-12", "Run full evaluation, compute before-after metrics, and perform manual spot-checking."),
            ("13-14", "Write final report, polish prototype, prepare demo and resume-ready project summary."),
        ],
        [Inches(1.1), Inches(5.4)],
    )

    add_heading(doc, "Expected Deliverables", 1)
    add_bullets(
        doc,
        [
            "A working prototype that orchestrates generation, execution, smell detection, feedback, and repair.",
            "A small Java benchmark or task set with logged initial tests, repaired tests, AromaDr reports, and execution results.",
            "An empirical evaluation comparing one-shot tests and repaired tests.",
            "A final written report explaining research questions, method, results, limitations, and future work.",
            "A demo-ready repository suitable for a resume or graduate research discussion.",
        ],
    )

    add_heading(doc, "Expected Contribution", 1)
    doc.add_paragraph(
        "The central contribution is a goal-driven repair framework for agent-"
        "generated Java tests. Existing coding agents often generate tests in a "
        "one-shot interaction, and existing test smell work often detects or "
        "characterizes quality issues after the fact. This project combines those "
        "directions by using AromaDr test-smell feedback as an active signal for "
        "iterative improvement. The project asks not only what kinds of Java tests "
        "agents generate, but whether a criteria-driven master agent can help "
        "coding agents produce tests that are more meaningful, maintainable, and "
        "useful."
    )

    add_heading(doc, "Short Pitch for Faculty Discussion", 1)
    doc.add_paragraph(
        "I am interested in building a goal-driven repair loop for agent-generated "
        "Java tests. A coding agent first generates JUnit tests, but a master agent "
        "then checks those tests against explicit quality criteria such as build "
        "success, passing status, assertion quality, mock usage, edge-case "
        "coverage, and AromaDr-detected test smells. If the tests do not satisfy "
        "the criteria, the master agent gives targeted smell-based feedback and "
        "asks the coding agent to revise them. I want to evaluate whether this "
        "iterative process reduces test smells and improves the usefulness of "
        "agent-generated Java tests compared with one-shot generation."
    )

    add_heading(doc, "Collaboration and Update Plan", 1)
    doc.add_paragraph(
        "I will use the #apl-alan Slack channel for project communication and can "
        "post concise weekly updates that summarize completed work, current "
        "blockers, next steps, and any questions for Professor Fazzini. If helpful, "
        "automated weekly Slack reminders can be used to make the update cadence "
        "consistent."
    )

    add_heading(doc, "References", 1)
    refs = [
        "Milanese, R., Salzano, F., Spina, A., Vitale, A., Pareschi, R., Fasano, F., and Fazzini, M. Human-Agent versus Human Pull Requests: A Testing-Focused Characterization and Comparison. MSR 2026 Mining Challenge. https://arxiv.org/abs/2601.21194",
        "AromaDr: A Language-Independent Tool for Detecting Test Smells.",
        "Van Deursen, A., Moonen, L., van den Bergh, A., and Kok, G. Refactoring Test Code. XP 2001. https://testsmells.org/",
        "Peruma, A. et al. Software Unit Test Smells and tsDetect resources. https://testsmells.org/",
        "Spadini, D. et al. Test smells 20 years later: detectability, validity, and reliability. Empirical Software Engineering. https://link.springer.com/article/10.1007/s10664-022-10207-5",
    ]
    for ref in refs:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        p.add_run(ref)

    doc.save(OUT)


if __name__ == "__main__":
    main()
